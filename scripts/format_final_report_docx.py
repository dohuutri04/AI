from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
IN_FILE = ROOT / "TEMPT_EduConnect_Full_v3.docx"
OUT_FILE = ROOT / "TEMPT_EduConnect_FINAL.docx"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def main():
    doc = Document(IN_FILE)

    # Global font defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    # Page setup: A4-ish margins for thesis
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.orientation = WD_SECTION.CONTINUOUS

    # Paragraph formatting: line spacing 1.5, justify body text
    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.line_spacing = 1.5
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        # Keep centered cover/title paragraphs as-is
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        # Keep list formatting alignment unchanged
        if "List" in style_name:
            continue
        # Keep heading alignment unchanged
        if "Heading" in style_name:
            continue
        if text:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)

    # Table content formatting
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

    # Footer page number
    for section in doc.sections:
        footer = section.footer
        if len(footer.paragraphs) == 0:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]
            para.clear()
        add_page_number(para)

    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
