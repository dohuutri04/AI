from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
IN_FILE = ROOT / "TEMPT_EduConnect_Full_v2.docx"
OUT_FILE = ROOT / "TEMPT_EduConnect_Full_v3.docx"


def should_append_period(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if t.endswith((".", "!", "?", ":", ";", "…")):
        return False
    # Skip short section labels/headers-like lines.
    if len(t.split()) <= 3:
        return False
    return True


def normalize_doc():
    doc = Document(IN_FILE)

    # Base style.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    # Ensure all existing runs use Times New Roman.
    for para in doc.paragraphs:
        text = para.text.strip()
        if should_append_period(text):
            para.add_run(".")
        for run in para.runs:
            run.font.name = "Times New Roman"

    # Apply for all table cells as well.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if should_append_period(text):
                        para.add_run(".")
                    for run in para.runs:
                        run.font.name = "Times New Roman"

    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    normalize_doc()
