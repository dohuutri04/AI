from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "TEMPT_EduConnect_Full_v2.docx"


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text):
    return doc.add_paragraph(text)


def b(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def n(doc, text):
    doc.add_paragraph(text, style="List Number")


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, x in enumerate(headers):
        t.rows[0].cells[i].text = str(x)
    for row in rows:
        c = t.add_row().cells
        for i, x in enumerate(row):
            c[i].text = str(x)


def pb(doc):
    doc.add_page_break()


def cover(doc):
    for line in [
        "TRUONG DAI HOC .............................................................",
        "KHOA ......................................................................",
        "BO MON ....................................................................",
    ]:
        x = p(doc, line)
        x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p(doc, "")
    x = p(doc, "BAO CAO DO AN TOT NGHIEP")
    x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(20)
    p(doc, "")
    x = p(doc, "DE TAI")
    x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    x.runs[0].bold = True
    x = p(doc, "HE THONG QUAN LY HOC TAP THONG MINH EDUCONNECT")
    x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(16)
    p(doc, "")
    p(doc, "Sinh vien thuc hien: ............................................................")
    p(doc, "MSSV: ...........................................................................")
    p(doc, "Lop: .............................................................................")
    p(doc, "Giang vien huong dan: ............................................................")
    p(doc, "Nien khoa: 2025 - 2026")
    p(doc, "")
    x = p(doc, "Binh Duong, 2026")
    x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def chapter_1(doc):
    h(doc, "CHUONG 1. TONG QUAN DE TAI", 1)
    h(doc, "1.1 Dat van de", 2)
    p(
        doc,
        "Dao tao truc tuyen da tro thanh xu huong tat yeu, nhung ti le bo hoc o cac khoa hoc online van cao "
        "do nguoi hoc thieu dong luc, thieu lo trinh ro rang va thieu su theo doi lien tuc. De tai huong den "
        "xay dung mot he thong LMS thong minh co kha nang can thiep som dua tren du lieu hanh vi.",
    )
    h(doc, "1.2 Ly do chon de tai", 2)
    for x in [
        "Nhu cau ca nhan hoa hoc tap ngay cang tang.",
        "Can mot he thong vua co AI vua tiet kiem chi phi van hanh.",
        "Can bo cong cu quan tri dong bo cho hoc vien, giang vien va admin.",
        "Can co che fallback de dam bao tinh san sang cua he thong.",
    ]:
        b(doc, x)
    h(doc, "1.3 Muc tieu", 2)
    for x in [
        "Xay dung nen tang hoc truc tuyen day du quy trinh.",
        "Theo doi tien do theo tung lesson va tung khoa hoc.",
        "Tinh risk score canh bao nguy co bo hoc.",
        "Goi y bai hoc tiep theo va khoa hoc de xuat.",
        "Trien khai AI hybrid: cloud mode + internal mode.",
    ]:
        n(doc, x)
    h(doc, "1.4 Doi tuong va pham vi", 2)
    p(
        doc,
        "Doi tuong nghien cuu la hanh vi hoc tap cua nguoi dung tren web app EduConnect. "
        "Pham vi bao gom he thong web monolithic; khong bao gom ung dung mobile native va thanh toan that.",
    )
    h(doc, "1.5 Phuong phap nghien cuu", 2)
    for x in [
        "Khao sat quy trinh hoc online va cac diem dung bo hoc.",
        "Phan tich yeu cau chuc nang/phi chuc nang.",
        "Thiet ke CSDL va API theo use case.",
        "Xay dung prototype, kiem thu, danh gia va cai tien.",
    ]:
        b(doc, x)
    h(doc, "1.6 Cau truc bao cao", 2)
    for x in [
        "Chuong 1: Tong quan",
        "Chuong 2: Co so ly thuyet va cong nghe",
        "Chuong 3: Phan tich thiet ke",
        "Chuong 4: Trien khai",
        "Chuong 5: Kiem thu danh gia",
        "Chuong 6: Ket luan va huong phat trien",
    ]:
        b(doc, x)


def chapter_2(doc):
    h(doc, "CHUONG 2. CO SO LY THUYET VA CONG NGHE", 1)
    h(doc, "2.1 He thong quan ly hoc tap (LMS)", 2)
    p(
        doc,
        "LMS la he thong quan ly noi dung hoc tap, quan ly nguoi hoc, giao bai, cham tien do va tong hop bao cao. "
        "Trong de tai nay, LMS duoc mo rong them lop AI personalization.",
    )
    h(doc, "2.2 Co so ly thuyet ve recommendation", 2)
    b(doc, "Content-based recommendation: dua tren muc do, chu de, lich su hoc.")
    b(doc, "Behavior-based recommendation: dua tren tien do, inactivity, completion.")
    b(doc, "Hybrid recommendation: ket hop nhieu tin hieu de nang do tin cay.")
    h(doc, "2.3 Co so ly thuyet ve risk scoring", 2)
    p(
        doc,
        "Risk score duoc tinh tu cac chi so: progress trung binh, ti le stalled, inactivity days. "
        "Diem tong hop duoc map sang cac muc low/medium/high de kich hoat chien luoc can thiep.",
    )
    h(doc, "2.4 Cong nghe su dung", 2)
    tbl(
        doc,
        ["Nhom", "Cong nghe", "Vai tro", "Ly do chon"],
        [
            ["Backend", "Flask", "Xu ly route + business", "Nhe, de hoc, phu hop do an"],
            ["Database", "SQLite", "Luu du lieu nghiep vu", "Khong can server rieng"],
            ["Frontend", "HTML/CSS/JS", "UI/UX", "Toi uu va de tuy bien"],
            ["AI cloud", "Gemini API", "Coach message nang cao", "Nhanh, de tich hop"],
            ["AI local", "Rule engine", "Fallback mien phi", "On dinh, tiet kiem chi phi"],
        ],
    )
    h(doc, "2.5 Ket luan chuong", 2)
    p(doc, "Chuong nay xac lap nen tang ly thuyet va bo cong nghe phuc vu cho cac chuong phan tich va trien khai.")


def chapter_3(doc):
    h(doc, "CHUONG 3. PHAN TICH VA THIET KE HE THONG", 1)
    h(doc, "3.1 Mo ta actor va use case", 2)
    tbl(
        doc,
        ["Actor", "Muc tieu", "Use case tieu bieu"],
        [
            ["Khach", "Kham pha he thong", "Xem trang chu, tim khoa hoc"],
            ["Hoc vien", "Hoc tap va theo doi tien do", "Enroll, hoc lesson, nhan AI coach"],
            ["Giang vien", "Tao noi dung dao tao", "Tao/sua/xoa khoa hoc, bai hoc"],
            ["Admin", "Giam sat va duyet", "Quan ly users, duyet giao dich, dashboard"],
        ],
    )
    h(doc, "3.2 Dac ta yeu cau chuc nang", 2)
    features = [
        "FR01 Dang ky/Dang nhap/ Dang xuat",
        "FR02 Quan ly profile ca nhan",
        "FR03 Tim kiem va loc khoa hoc",
        "FR04 Dang ky khoa hoc (co tinh phi tu vi)",
        "FR05 Hoc bai hoc va danh dau hoan thanh",
        "FR06 Goi y AI: risk, next lesson, recommendations",
        "FR07 Quan ly khoa hoc cho giang vien",
        "FR08 Wallet: nap/rut request va lich su",
        "FR09 Admin duyet nap/rut, tu choi rut co refund",
    ]
    for f in features:
        b(doc, f)
    h(doc, "3.3 Dac ta yeu cau phi chuc nang", 2)
    for nfr in [
        "NFR01 Do dung: giao dien ro rang, thao tac don gian.",
        "NFR02 Bao mat: session + phan quyen route.",
        "NFR03 Tin cay: AI fallback khi cloud loi.",
        "NFR04 Bao tri: tach service de de mo rong.",
        "NFR05 Hieu nang: phan hoi API o muc chap nhan duoc.",
    ]:
        b(doc, nfr)
    h(doc, "3.4 Thiet ke CSDL", 2)
    tbl(
        doc,
        ["Bang", "Khoa chinh", "Quan he", "Y nghia"],
        [
            ["users", "id", "1-n enrollments", "Thong tin tai khoan va vi"],
            ["courses", "id", "1-n lessons", "Thong tin khoa hoc"],
            ["lessons", "id", "n-1 courses", "Noi dung bai hoc"],
            ["enrollments", "id", "n-1 users/courses", "Dang ky hoc va progress"],
            ["lesson_progress", "(user_id,lesson_id)", "n-1 users/lessons", "Hoan thanh theo bai"],
            ["wallet_transactions", "id", "n-1 users", "Lich su nap/rut/mua/refund"],
        ],
    )
    h(doc, "3.5 Thiet ke AI personalization", 2)
    for step in [
        "B1: Lay du lieu enrollments + lesson_progress",
        "B2: Tinh avg_progress, stalled_count, inactivity_days",
        "B3: Tinh risk_score va risk_level",
        "B4: Tao danh sach next_lessons",
        "B5: Tao recommendations theo category/level",
        "B6: Sinh coach message (gemini hoac internal)",
    ]:
        n(doc, step)
    h(doc, "3.6 Thiet ke giao dien", 2)
    p(
        doc,
        "Giao dien duoc thiet ke theo phong cach dashboard hien dai, dark-light contrast. "
        "Tab AI co cac card thong tin: segment, risk, inactivity, coach message, next lessons, recommendations.",
    )
    h(doc, "3.7 Ket luan chuong", 2)
    p(doc, "He thong duoc thiet ke day du o ca 3 lop: presentation, business, data.")


def chapter_4(doc):
    h(doc, "CHUONG 4. TRIEN KHAI HE THONG", 1)
    h(doc, "4.1 Moi truong phat trien", 2)
    b(doc, "OS: Windows 10")
    b(doc, "Python 3.10+")
    b(doc, "Flask + sqlite3 + python-docx")
    b(doc, "Trinh duyet: Chrome/CocCoc")
    h(doc, "4.2 Cau truc ma nguon", 2)
    b(doc, "app.py: route layer + integration")
    b(doc, "services/ai_personalization.py: AI service")
    b(doc, "templates/: giao dien Jinja2")
    b(doc, "static/js, static/css: xu ly client")
    b(doc, "database.sql: schema + seed")
    h(doc, "4.3 Trien khai API tieu bieu", 2)
    tbl(
        doc,
        ["API", "Method", "Input", "Output"],
        [
            ["/api/ai/personalization", "GET", "session user_id", "profile + coach"],
            ["/enroll/<id>", "POST", "course_id", "success + first_lesson"],
            ["/mark-lesson-complete", "POST", "lesson_id", "progress_pct"],
            ["/wallet/withdraw-request", "POST", "amount", "pending withdrawal"],
            ["/admin/withdrawals/reject/<id>", "POST", "note", "refund transaction"],
        ],
    )
    h(doc, "4.4 Trien khai AI mode", 2)
    p(doc, "He thong ho tro bien moi truong AI_COACH_MODE:")
    b(doc, "auto: uu tien cloud AI neu kha dung.")
    b(doc, "internal: luon dung AI mien phi noi bo.")
    p(
        doc,
        "UI hien thi ro mode dang su dung: AI Coach (Gemini Mode) hoac AI Coach (Internal/Free Mode).",
    )
    h(doc, "4.5 Trien khai chuc nang wallet va refund", 2)
    p(
        doc,
        "Khi user gui yeu cau rut tien, so du duoc tam tru va tao giao dich pending. "
        "Neu admin tu choi, he thong cong lai so du vao vi, danh dau giao dich rut la failed "
        "va tao them giao dich refund de minh bach.",
    )
    h(doc, "4.6 Ket luan chuong", 2)
    p(doc, "Qua trinh trien khai dat muc tieu de tai, he thong van hanh on dinh va de theo doi.")


def chapter_5(doc):
    h(doc, "CHUONG 5. KIEM THU VA DANH GIA", 1)
    h(doc, "5.1 Ke hoach kiem thu", 2)
    p(doc, "Kiem thu duoc thuc hien theo nhom: auth, hoc tap, AI, wallet, admin.")
    h(doc, "5.2 Bang test case chi tiet", 2)
    tbl(
        doc,
        ["ID", "Nhom", "Tien dieu kien", "Buoc test", "Ket qua mong doi"],
        [
            ["TC01", "Auth", "Chua dang nhap", "Dang ky tai khoan moi", "Dang ky thanh cong"],
            ["TC02", "Auth", "Da co tai khoan", "Dang nhap", "Tao session user"],
            ["TC03", "Learning", "Da enroll", "Mo bai hoc + mark complete", "Tang tien do"],
            ["TC04", "AI", "AI mode internal", "GET /api/ai/personalization", "coach_source=rule_engine"],
            ["TC05", "AI", "AI mode auto + cloud loi", "Goi API", "Fallback noi bo"],
            ["TC06", "Wallet", "So du du", "Gui yeu cau rut", "Tao pending transaction"],
            ["TC07", "Admin", "Co pending withdrawal", "Reject withdrawal", "Refund + failed status"],
            ["TC08", "Admin", "Da reject", "Reject lan 2", "Bao loi khong con pending"],
        ],
    )
    h(doc, "5.3 Ket qua thuc nghiem", 2)
    p(
        doc,
        "Tat ca luong chinh deu dat ket qua mong doi. He thong giu duoc tinh lien tuc nho fallback AI noi bo, "
        "khong phu thuoc hoan toan vao cloud quota.",
    )
    h(doc, "5.4 Danh gia uu nhuoc diem", 2)
    b(doc, "Uu diem: de trien khai, de demo, de bao tri, AI mode linh hoat.")
    b(doc, "Han che: chua co bo test tu dong day du, chua toi uu cho tai luong lon.")
    h(doc, "5.5 Chi so danh gia", 2)
    tbl(
        doc,
        ["Chi so", "Gia tri hien tai", "Muc tieu tiep theo"],
        [
            ["Ti le test case dat", ">= 90%", ">= 95%"],
            ["Do san sang AI", "High (co fallback)", "High + cache"],
            ["Thoi gian phan hoi API", "Dat muc demo", "Toi uu hon voi cache"],
            ["Minh bach giao dich", "Co lich su refund", "Them audit log chi tiet"],
        ],
    )


def chapter_6(doc):
    h(doc, "CHUONG 6. KET LUAN VA HUONG PHAT TRIEN", 1)
    h(doc, "6.1 Ket luan", 2)
    p(
        doc,
        "De tai da dat duoc muc tieu xay dung he thong quan ly hoc tap thong minh co AI personalization. "
        "Giai phap hybrid AI giup can bang giua chat luong va chi phi van hanh.",
    )
    h(doc, "6.2 Dong gop chinh cua de tai", 2)
    for x in [
        "Xay dung duoc mot LMS day du nghiep vu cot loi.",
        "Trien khai thanh cong AI risk scoring va recommendation.",
        "Thiet ke fallback internal mode de khong phat sinh chi phi cloud bat buoc.",
        "Bo sung luong refund ro rang cho giao dich rut bi tu choi.",
    ]:
        b(doc, x)
    h(doc, "6.3 Huong phat trien", 2)
    for x in [
        "Tich hop bo test tu dong CI/CD.",
        "Them dashboard analytics nang cao cho giang vien.",
        "Adaptive learning sequencing theo ket qua bai tap.",
        "Toi uu performance bang cache + async jobs.",
        "Tich hop mobile app trong giai doan tiep theo.",
    ]:
        n(doc, x)


def appendices(doc):
    h(doc, "PHU LUC", 1)
    h(doc, "Phu luc A. Danh sach endpoint chinh", 2)
    endpoints = [
        ("GET", "/", "Trang chu"),
        ("GET", "/khoa-hoc", "Danh sach khoa hoc"),
        ("POST", "/enroll/<id>", "Dang ky khoa hoc"),
        ("GET", "/xem-bai-hoc/<id>", "Hoc bai hoc"),
        ("POST", "/mark-lesson-complete", "Danh dau hoan thanh"),
        ("GET", "/api/ai/personalization", "Lay goi y AI"),
        ("POST", "/wallet/deposit-request", "Gui yeu cau nap"),
        ("POST", "/wallet/withdraw-request", "Gui yeu cau rut"),
        ("POST", "/admin/withdrawals/complete/<id>", "Xac nhan da chuyen"),
        ("POST", "/admin/withdrawals/reject/<id>", "Tu choi va hoan tien"),
    ]
    tbl(doc, ["Method", "Endpoint", "Mo ta"], endpoints)

    h(doc, "Phu luc B. Danh sach bang du lieu", 2)
    for x in [
        "users, categories, courses, lessons",
        "lesson_materials, lesson_exercises",
        "enrollments, lesson_progress",
        "wallet_transactions, deposit_requests",
        "contacts, password_resets, delete_requests",
    ]:
        b(doc, x)

    h(doc, "Phu luc C. Danh sach so do da xay dung", 2)
    for x in [
        "docs/so-do-usecase.mmd",
        "docs/so-do-kien-truc-he-thong.mmd",
        "docs/so-do-er-rut-gon.mmd",
        "docs/so-do-luong-ai-ca-nhan-hoa.mmd",
    ]:
        b(doc, x)

    h(doc, "Phu luc D. Ke hoach trien khai thuc te", 2)
    tbl(
        doc,
        ["Giai doan", "Noi dung", "Thoi gian du kien"],
        [
            ["GD1", "Hoan thien product MVP", "2-4 tuan"],
            ["GD2", "Pilot voi nhom nguoi dung nho", "2 tuan"],
            ["GD3", "Toi uu AI + bo test tu dong", "4-6 tuan"],
            ["GD4", "Mo rong tich hop mobile", "6-8 tuan"],
        ],
    )


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    cover(doc)
    pb(doc)
    h(doc, "LOI CAM ON", 1)
    p(
        doc,
        "Em xin chan thanh cam on quy Thay/Cô va khoa Cong nghe thong tin da ho tro em trong suot qua trinh hoc tap va thuc hien do an.",
    )
    h(doc, "TOM TAT", 1)
    p(
        doc,
        "Bao cao trinh bay viec xay dung he thong LMS thong minh EduConnect voi AI personalization, bao gom recommendation, risk scoring va coach message hybrid.",
    )
    pb(doc)
    h(doc, "MUC LUC TOM TAT", 1)
    for x in [
        "Chuong 1 Tong quan",
        "Chuong 2 Co so ly thuyet",
        "Chuong 3 Phan tich thiet ke",
        "Chuong 4 Trien khai",
        "Chuong 5 Kiem thu danh gia",
        "Chuong 6 Ket luan",
        "Phu luc",
    ]:
        n(doc, x)
    pb(doc)
    chapter_1(doc)
    pb(doc)
    chapter_2(doc)
    pb(doc)
    chapter_3(doc)
    pb(doc)
    chapter_4(doc)
    pb(doc)
    chapter_5(doc)
    pb(doc)
    chapter_6(doc)
    pb(doc)
    appendices(doc)

    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
