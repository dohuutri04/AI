from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "TEMPT.docx"


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    doc.add_heading("BAO CAO DE TAI DO AN NGANH", level=0)
    doc.add_paragraph("Ten de tai: He thong quan ly hoc tap thong minh EduConnect")
    doc.add_paragraph("Sinh vien thuc hien: ..........................................")
    doc.add_paragraph("MSSV: ........................................................")
    doc.add_paragraph("Giang vien huong dan: ........................................")

    doc.add_heading("1. Gioi thieu de tai", level=1)
    doc.add_paragraph(
        "EduConnect la nen tang hoc truc tuyen xay dung tren Flask + SQLite, "
        "tap trung vao ca nhan hoa qua trinh hoc bang AI dua tren hanh vi nguoi dung."
    )

    doc.add_heading("2. Muc tieu chinh", level=1)
    add_bullet(doc, "Xay dung he thong quan ly hoc truc tuyen day du vai tro.")
    add_bullet(doc, "Theo doi tien do hoc tap theo tung bai hoc va tung khoa hoc.")
    add_bullet(doc, "Ca nhan hoa lo trinh hoc bang AI dua tren du lieu hanh vi.")
    add_bullet(doc, "Toi uu van hanh voi AI hybrid: cloud neu co, fallback noi bo neu khong.")

    doc.add_heading("3. Chuc nang trong tam", level=1)
    add_bullet(doc, "Dang ky, dang nhap, quan ly tai khoan va thong tin nguoi dung.")
    add_bullet(doc, "Tim kiem, loc khoa hoc, dang ky hoc va hoc bai hoc truc tuyen.")
    add_bullet(doc, "Danh dau hoan thanh bai hoc, theo doi tien do hoc tap.")
    add_bullet(doc, "Quan ly khoa hoc cua giang vien.")
    add_bullet(doc, "Quan tri he thong cho admin: nguoi dung, danh muc, lien he, giao dich.")

    doc.add_heading("4. Chuc nang AI cua de tai", level=1)
    add_bullet(doc, "Phan tich profile hoc vien: progress, inactivity, completed/stalled.")
    add_bullet(doc, "Tinh diem rui ro bo hoc (risk score) va muc rui ro.")
    add_bullet(doc, "Goi y bai hoc tiep theo dua tren trang thai hoc hien tai.")
    add_bullet(doc, "Goi y khoa hoc ca nhan hoa theo muc do va chu de quan tam.")
    add_bullet(doc, "AI Coach phan hoi theo 2 che do:")
    add_bullet(doc, " - Gemini (neu co quota/key).")
    add_bullet(doc, " - Internal/Free mode (rule engine) khi khong co cloud AI.")

    doc.add_heading("5. Kien truc va cong nghe", level=1)
    add_bullet(doc, "Backend: Python Flask.")
    add_bullet(doc, "Database: SQLite.")
    add_bullet(doc, "Frontend: HTML, CSS, JavaScript.")
    add_bullet(doc, "AI: Gemini API + Rule-based fallback.")

    doc.add_heading("6. Ket qua dat duoc", level=1)
    add_bullet(doc, "He thong chay on dinh voi luong hoc truc tuyen day du.")
    add_bullet(doc, "AI ca nhan hoa hoat dong duoc ca khi khong ton phi cloud.")
    add_bullet(doc, "Co co che auto-refresh AI va hien thi mode ro rang tren giao dien.")
    add_bullet(doc, "Ho tro xu ly giao dich vi va hoan tien khi tu choi rut.")

    doc.add_heading("7. Huong phat trien", level=1)
    add_bullet(doc, "Bo sung test tu dong cho cac luong quan trong.")
    add_bullet(doc, "Nang cap mo hinh du doan ket qua hoc tap va adaptive learning.")
    add_bullet(doc, "Tich hop thong bao email/app theo su kien.")

    doc.save(OUT_FILE)
    print(f"Updated: {OUT_FILE}")


if __name__ == "__main__":
    main()
