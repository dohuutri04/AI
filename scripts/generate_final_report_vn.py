from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "TEMPT_EduConnect_FINAL.docx"


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text):
    return doc.add_paragraph(text)


def b(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def n(doc, text):
    doc.add_paragraph(text, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, x in enumerate(headers):
        t.rows[0].cells[i].text = x
    for row in rows:
        c = t.add_row().cells
        for i, x in enumerate(row):
            c[i].text = str(x)


def page_break(doc):
    doc.add_page_break()


def format_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)

    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.line_spacing = 1.5
        style_name = para.style.name if para.style else ""
        if "Heading" not in style_name and "List" not in style_name:
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)


def cover(doc):
    x = p(doc, "TRƯỜNG ĐẠI HỌC .............................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x = p(doc, "KHOA ............................................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x = p(doc, "BỘ MÔN ..........................................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "")
    x = p(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(20)
    p(doc, "")
    x = p(doc, "ĐỀ TÀI")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x = p(doc, "HỆ THỐNG QUẢN LÝ HỌC TẬP THÔNG MINH EDUCONNECT")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(16)
    p(doc, "")
    p(doc, "Sinh viên thực hiện: ............................................................")
    p(doc, "MSSV: ...........................................................................")
    p(doc, "Lớp: .............................................................................")
    p(doc, "Giảng viên hướng dẫn: ............................................................")
    p(doc, "Ngành: Công nghệ thông tin")
    p(doc, "")
    x = p(doc, "Bình Dương, năm 2026")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER


def body(doc):
    page_break(doc)
    h(doc, "LỜI CẢM ƠN", 1)
    p(
        doc,
        "Em xin chân thành cảm ơn quý Thầy/Cô khoa Công nghệ thông tin đã hỗ trợ em trong suốt quá trình học tập và thực hiện đồ án. "
        "Đặc biệt, em xin gửi lời cảm ơn đến giảng viên hướng dẫn đã tận tình góp ý để em hoàn thiện đề tài này.",
    )

    h(doc, "TÓM TẮT ĐỀ TÀI", 1)
    p(
        doc,
        "Đề tài xây dựng hệ thống quản lý học tập trực tuyến EduConnect có tích hợp AI cá nhân hóa. "
        "Hệ thống hỗ trợ theo dõi tiến độ, chấm điểm rủi ro bỏ học, gợi ý bài học tiếp theo, gợi ý khóa học phù hợp và phản hồi AI Coach theo chế độ Hybrid.",
    )

    page_break(doc)
    h(doc, "MỤC LỤC TÓM TẮT", 1)
    for x in [
        "Chương 1. Tổng quan đề tài",
        "Chương 2. Cơ sở lý thuyết và công nghệ",
        "Chương 3. Phân tích và thiết kế hệ thống",
        "Chương 4. Triển khai hệ thống",
        "Chương 5. Kiểm thử và đánh giá",
        "Chương 6. Kết luận và hướng phát triển",
        "Phụ lục",
    ]:
        n(doc, x)

    page_break(doc)
    h(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    h(doc, "1.1 Đặt vấn đề", 2)
    p(
        doc,
        "Trong bối cảnh chuyển đổi số giáo dục, nhu cầu học trực tuyến tăng nhanh nhưng tỷ lệ bỏ học vẫn cao do thiếu cá nhân hóa. "
        "Đề tài hướng đến giải pháp LMS thông minh có khả năng phân tích dữ liệu hành vi và đưa ra gợi ý học tập phù hợp.",
    )
    h(doc, "1.2 Mục tiêu", 2)
    for x in [
        "Xây dựng nền tảng LMS đầy đủ cho học viên, giảng viên và quản trị viên.",
        "Theo dõi tiến độ học tập theo từng bài học và khóa học.",
        "Chấm điểm rủi ro bỏ học dựa trên dữ liệu hành vi.",
        "Gợi ý bài học kế tiếp và khóa học phù hợp.",
        "Triển khai cơ chế AI Hybrid: cloud mode và internal/free mode.",
    ]:
        b(doc, x)

    h(doc, "1.3 Phạm vi", 2)
    p(
        doc,
        "Đề tài triển khai trên nền tảng web, kiến trúc monolithic, tập trung vào chức năng cốt lõi và khả năng demo thực tế. "
        "Các tính năng mobile native và thanh toán thực không nằm trong phạm vi bản hiện tại.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", 1)
    h(doc, "2.1 Nền tảng lý thuyết", 2)
    b(doc, "Mô hình Client-Server cho ứng dụng web.")
    b(doc, "Nguyên lý Recommendation dựa trên hành vi người dùng.")
    b(doc, "Risk Scoring cho bài toán dự đoán nguy cơ bỏ học.")
    b(doc, "Thiết kế hệ thống có khả năng fallback khi cloud AI lỗi.")

    h(doc, "2.2 Công nghệ sử dụng", 2)
    table(
        doc,
        ["Thành phần", "Công nghệ", "Mục đích", "Lý do chọn"],
        [
            ["Backend", "Python Flask", "Xử lý route và nghiệp vụ", "Nhẹ, nhanh, dễ mở rộng"],
            ["Database", "SQLite", "Lưu trữ dữ liệu hệ thống", "Đơn giản, không cần server riêng"],
            ["Frontend", "HTML/CSS/JS", "Xây dựng giao diện", "Linh hoạt, dễ tùy chỉnh"],
            ["AI Cloud", "Gemini API", "Sinh phản hồi AI nâng cao", "Tích hợp nhanh"],
            ["AI Nội bộ", "Rule-based", "Fallback miễn phí", "Ổn định, không phụ thuộc quota"],
        ],
    )

    page_break(doc)
    h(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    h(doc, "3.1 Actor và Use Case", 2)
    table(
        doc,
        ["Actor", "Vai trò", "Use Case chính"],
        [
            ["Khách", "Người chưa đăng nhập", "Xem trang chủ, tìm kiếm khóa học"],
            ["Học viên", "Người học", "Đăng ký khóa, học bài, nhận gợi ý AI"],
            ["Giảng viên", "Người tạo nội dung", "Tạo/sửa/xóa khóa học và bài học"],
            ["Admin", "Quản trị", "Duyệt giao dịch, quản lý dữ liệu hệ thống"],
        ],
    )
    h(doc, "3.2 Đặc tả yêu cầu chức năng", 2)
    for x in [
        "FR01: Đăng ký/đăng nhập/đăng xuất.",
        "FR02: Tìm kiếm và lọc khóa học.",
        "FR03: Đăng ký khóa học và học bài học.",
        "FR04: Đánh dấu hoàn thành bài học.",
        "FR05: API AI cá nhân hóa.",
        "FR06: Ví điện tử và giao dịch nạp/rút.",
        "FR07: Admin từ chối rút và hoàn tiền tự động.",
    ]:
        b(doc, x)

    h(doc, "3.3 Thiết kế cơ sở dữ liệu", 2)
    table(
        doc,
        ["Bảng", "Mô tả"],
        [
            ["users", "Thông tin tài khoản, phân quyền, số dư ví"],
            ["courses", "Thông tin khóa học, giá, cấp độ, danh mục"],
            ["lessons", "Bài học thuộc từng khóa"],
            ["enrollments", "Quan hệ học viên - khóa học và tiến độ"],
            ["lesson_progress", "Dữ liệu hoàn thành bài học"],
            ["wallet_transactions", "Lịch sử nạp/rút/mua/hoàn tiền"],
        ],
    )

    h(doc, "3.4 Luồng AI cá nhân hóa", 2)
    for x in [
        "Thu thập dữ liệu hành vi học tập.",
        "Tính các chỉ số: avg_progress, stalled_count, inactivity_days.",
        "Tính risk_score và risk_level.",
        "Gợi ý bài học tiếp theo và khóa học đề xuất.",
        "Sinh coach message theo chế độ AI hiện tại.",
    ]:
        n(doc, x)

    page_break(doc)
    h(doc, "CHƯƠNG 4. TRIỂN KHAI HỆ THỐNG", 1)
    h(doc, "4.1 Môi trường phát triển", 2)
    for x in [
        "OS: Windows 10",
        "Python 3.10+",
        "Flask + sqlite3 + python-docx",
        "Trình duyệt: Chrome/Cốc Cốc",
    ]:
        b(doc, x)

    h(doc, "4.2 Cấu trúc mã nguồn", 2)
    for x in [
        "app.py: route layer + integration",
        "services/ai_personalization.py: AI service",
        "templates/: giao diện Jinja2",
        "static/js, static/css: xử lý client",
        "database.sql: schema + seed",
    ]:
        b(doc, x)

    h(doc, "4.3 Triển khai API tiêu biểu", 2)
    table(
        doc,
        ["API", "Method", "Input", "Output"],
        [
            ["/api/ai/personalization", "GET", "session user_id", "profile + coach"],
            ["/enroll/<id>", "POST", "course_id", "success + first_lesson"],
            ["/mark-lesson-complete", "POST", "lesson_id", "progress_pct"],
            ["/wallet/withdraw-request", "POST", "amount", "pending withdrawal"],
            ["/admin/withdrawals/reject/<id>", "POST", "note", "refund transaction"],
        ],
    )

    h(doc, "4.4 Cơ chế AI Hybrid", 2)
    p(
        doc,
        "Hệ thống hỗ trợ biến môi trường AI_COACH_MODE gồm hai chế độ: "
        "auto (ưu tiên cloud AI nếu khả dụng) và internal (luôn dùng AI nội bộ miễn phí). "
        "Cơ chế này đảm bảo hệ thống luôn hoạt động kể cả khi hết quota cloud.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    h(doc, "5.1 Kịch bản kiểm thử chính", 2)
    table(
        doc,
        ["ID", "Mô tả", "Kết quả mong đợi"],
        [
            ["TC01", "Đăng ký tài khoản mới", "Tạo user thành công"],
            ["TC02", "Đăng ký khóa học", "Thêm enrollment hợp lệ"],
            ["TC03", "Đánh dấu hoàn thành bài", "Cập nhật tiến độ"],
            ["TC04", "AI internal mode", "coach_source = rule_engine"],
            ["TC05", "AI auto mode khi cloud lỗi", "Fallback nội bộ"],
            ["TC06", "Admin từ chối rút", "Hoàn tiền + ghi refund"],
        ],
    )
    h(doc, "5.2 Đánh giá tổng quan", 2)
    p(
        doc,
        "Các chức năng cốt lõi đều đạt yêu cầu. Hệ thống AI nội bộ vận hành ổn định, "
        "phù hợp cho bối cảnh cần tối ưu chi phí nhưng vẫn đảm bảo trải nghiệm cá nhân hóa.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    h(doc, "6.1 Kết luận", 2)
    p(
        doc,
        "Đề tài đã xây dựng thành công hệ thống quản lý học tập thông minh EduConnect với các chức năng hoàn chỉnh và lớp AI cá nhân hóa. "
        "Mô hình AI Hybrid giúp cân bằng giữa chất lượng phản hồi và chi phí vận hành.",
    )
    h(doc, "6.2 Hướng phát triển", 2)
    for x in [
        "Tích hợp bộ kiểm thử tự động CI/CD.",
        "Nâng cấp adaptive learning theo kết quả bài tập.",
        "Bổ sung dashboard phân tích học tập nâng cao.",
        "Mở rộng tích hợp mobile app.",
    ]:
        n(doc, x)

    page_break(doc)
    h(doc, "PHỤ LỤC", 1)
    h(doc, "Phụ lục A. Danh sách endpoint tiêu biểu", 2)
    table(
        doc,
        ["Method", "Endpoint", "Mô tả"],
        [
            ["GET", "/", "Trang chủ"],
            ["GET", "/khoa-hoc", "Danh sách khóa học"],
            ["POST", "/enroll/<id>", "Đăng ký khóa học"],
            ["POST", "/mark-lesson-complete", "Đánh dấu hoàn thành bài học"],
            ["GET", "/api/ai/personalization", "Lấy gợi ý AI"],
            ["POST", "/admin/withdrawals/reject/<id>", "Từ chối rút và hoàn tiền"],
        ],
    )

    h(doc, "Phụ lục B. Danh sách sơ đồ", 2)
    for x in [
        "docs/so-do-usecase.mmd",
        "docs/so-do-kien-truc-he-thong.mmd",
        "docs/so-do-er-rut-gon.mmd",
        "docs/so-do-luong-ai-ca-nhan-hoa.mmd",
    ]:
        b(doc, x)


def main():
    doc = Document()
    cover(doc)
    body(doc)
    format_doc(doc)
    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
