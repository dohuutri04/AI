from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT_FILE = ROOT / "TEMPT_EduConnect_FINAL_CHI_TIET.docx"
TITLE = "EduConnect: Nền tảng học trực tuyến cá nhân hóa bằng AI"


def h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def p(doc: Document, text: str):
    return doc.add_paragraph(text)


def b(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def n(doc: Document, text: str):
    doc.add_paragraph(text, style="List Number")


def tbl(doc: Document, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, x in enumerate(headers):
        t.rows[0].cells[i].text = str(x)
    for row in rows:
        cells = t.add_row().cells
        for i, x in enumerate(row):
            cells[i].text = str(x)
    return t


def page_break(doc: Document):
    doc.add_page_break()


def add_code_block(doc: Document, text: str):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    return para


def add_word_field(paragraph, instr: str):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    paragraph._p.append(fld)
    return fld


def add_toc(paragraph):
    add_word_field(paragraph, r'TOC \o "1-3" \h \z \u')


def add_list_of_figures(paragraph):
    add_word_field(paragraph, r'TOC \h \z \c "Hình"')


def add_list_of_tables(paragraph):
    add_word_field(paragraph, r'TOC \h \z \c "Bảng"')


def review_pages(doc: Document):
    page_break(doc)
    h(doc, "PHIẾU NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
    p(doc, "Tên đề tài: " + TITLE)
    p(doc, "Sinh viên thực hiện: ............................................................")
    p(doc, "MSSV: ...........................................................................")
    p(doc, "Giảng viên hướng dẫn: ............................................................")
    p(doc, "")
    p(doc, "1. Nội dung nhận xét:")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "")
    p(doc, "2. Đánh giá:")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "")
    p(doc, "3. Kết luận:   ☐ Đồng ý cho bảo vệ      ☐ Không đồng ý cho bảo vệ")
    p(doc, "")
    p(doc, "Bình Dương, ngày ..... tháng ..... năm 2026")
    p(doc, "Giảng viên hướng dẫn (Ký, ghi rõ họ tên)")

    page_break(doc)
    h(doc, "PHIẾU NHẬN XÉT CỦA GIẢNG VIÊN PHẢN BIỆN", 1)
    p(doc, "Tên đề tài: " + TITLE)
    p(doc, "Sinh viên thực hiện: ............................................................")
    p(doc, "MSSV: ...........................................................................")
    p(doc, "Giảng viên phản biện: ...........................................................")
    p(doc, "")
    p(doc, "1. Nội dung nhận xét:")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "")
    p(doc, "2. Đánh giá:")
    p(doc, "........................................................................................................")
    p(doc, "........................................................................................................")
    p(doc, "")
    p(doc, "3. Kết luận:   ☐ Đồng ý cho bảo vệ      ☐ Không đồng ý cho bảo vệ")
    p(doc, "")
    p(doc, "Bình Dương, ngày ..... tháng ..... năm 2026")
    p(doc, "Giảng viên phản biện (Ký, ghi rõ họ tên)")


def apply_global_format(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        pf = para.paragraph_format
        pf.line_spacing = 1.5
        if "Heading" not in style_name and "List" not in style_name:
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER and para.text.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Times New Roman"
            if not run.font.size:
                run.font.size = Pt(13)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.3
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)


def cover(doc: Document):
    x = p(doc, "TRƯỜNG ĐẠI HỌC .............................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x = p(doc, "KHOA ............................................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x = p(doc, "BỘ MÔN ..........................................................................")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "")
    x = p(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(20)
    p(doc, "")
    x = p(doc, "TÊN ĐỀ TÀI")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x = p(doc, TITLE)
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.runs[0].bold = True
    x.runs[0].font.size = Pt(16)
    p(doc, "")
    p(doc, "Sinh viên thực hiện: ............................................................")
    p(doc, "MSSV: ...........................................................................")
    p(doc, "Lớp: .............................................................................")
    p(doc, "Giảng viên hướng dẫn: ............................................................")
    p(doc, "Ngành: Công nghệ thông tin")
    p(doc, "")
    x = p(doc, "Bình Dương, năm 2026")
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER


def content(doc: Document):
    page_break(doc)
    h(doc, "LỜI CẢM ƠN", 1)
    p(
        doc,
        "Em xin chân thành cảm ơn quý Thầy/Cô khoa Công nghệ thông tin đã hỗ trợ em trong quá trình học tập và thực hiện đồ án. "
        "Đặc biệt, em xin gửi lời cảm ơn sâu sắc đến giảng viên hướng dẫn đã tận tình định hướng, góp ý và đồng hành để em hoàn thiện đề tài.",
    )

    h(doc, "TÓM TẮT ĐỀ TÀI", 1)
    p(
        doc,
        f"Đề tài \"{TITLE}\" xây dựng một hệ thống học trực tuyến có khả năng cá nhân hóa dựa trên dữ liệu hành vi học tập. "
        "Hệ thống hỗ trợ quản lý người dùng, khóa học, bài học, theo dõi tiến độ theo từng bài và tích hợp AI Coach để gợi ý lộ trình học phù hợp.",
    )
    p(
        doc,
        "Giải pháp AI được thiết kế theo mô hình Hybrid: ưu tiên cloud AI khi khả dụng, và tự động fallback sang AI nội bộ miễn phí khi không có quota hoặc dịch vụ ngoài lỗi.",
    )

    page_break(doc)
    h(doc, "MỤC LỤC", 1)
    p(doc, "Nhấn chuột phải vào mục lục và chọn Update Field để cập nhật tự động.")
    add_toc(doc.add_paragraph())

    page_break(doc)
    h(doc, "DANH MỤC HÌNH", 1)
    p(doc, "Nhấn chuột phải vào danh mục và chọn Update Field để cập nhật tự động.")
    add_list_of_figures(doc.add_paragraph())

    page_break(doc)
    h(doc, "DANH MỤC BẢNG", 1)
    p(doc, "Nhấn chuột phải vào danh mục và chọn Update Field để cập nhật tự động.")
    add_list_of_tables(doc.add_paragraph())

    page_break(doc)
    h(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    h(doc, "1.1 Bối cảnh và bài toán", 2)
    p(
        doc,
        "Trong môi trường học trực tuyến, người học thường bỏ dở khóa giữa chừng do không có định hướng rõ ràng và thiếu cơ chế theo dõi sát sao. "
        "Bài toán đặt ra là xây dựng một nền tảng LMS không chỉ quản lý nội dung học tập mà còn chủ động hỗ trợ người học bằng trí tuệ nhân tạo.",
    )
    h(doc, "1.2 Mục tiêu tổng quát", 2)
    for x in [
        "Xây dựng nền tảng học trực tuyến có đầy đủ vai trò: học viên, giảng viên, quản trị viên.",
        "Theo dõi tiến độ học tập theo từng bài học, khóa học và mốc hoàn thành.",
        "Chấm điểm rủi ro bỏ học dựa trên hành vi thực tế của người dùng.",
        "Gợi ý bài học tiếp theo và khóa học đề xuất theo năng lực hiện tại.",
        "Tạo cơ chế AI hoạt động ổn định, không phụ thuộc hoàn toàn vào dịch vụ trả phí.",
    ]:
        b(doc, x)
    h(doc, "1.3 Phạm vi và giới hạn", 2)
    p(
        doc,
        "Đề tài tập trung vào web app cho môi trường demo và học thuật. "
        "Các chức năng thanh toán thật, hạ tầng microservices quy mô lớn hoặc mobile native chưa nằm trong phạm vi phiên bản hiện tại.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", 1)
    h(doc, "2.1 Cơ sở lý thuyết", 2)
    b(doc, "Mô hình kiến trúc web client-server.")
    b(doc, "LMS và quy trình quản lý học tập số.")
    b(doc, "Recommendation theo hành vi và mức độ học.")
    b(doc, "Risk scoring để phát hiện sớm nguy cơ bỏ học.")
    h(doc, "2.2 Công nghệ sử dụng", 2)
    tbl(
        doc,
        ["Thành phần", "Công nghệ", "Vai trò", "Lý do lựa chọn"],
        [
            ["Backend", "Python Flask", "Route, nghiệp vụ, API", "Nhẹ, linh hoạt, phù hợp đồ án"],
            ["Database", "SQLite", "Lưu trữ dữ liệu hệ thống", "Đơn giản, dễ triển khai"],
            ["Frontend", "HTML/CSS/JS", "Giao diện và tương tác", "Dễ tùy chỉnh UI/UX"],
            ["AI Cloud", "Gemini API", "Sinh phản hồi coach nâng cao", "Tích hợp nhanh"],
            ["AI nội bộ", "Rule-based service", "Fallback miễn phí", "Ổn định và tiết kiệm"],
        ],
    )
    h(doc, "2.3 Mô hình AI Hybrid", 2)
    p(
        doc,
        "Hệ thống cho phép cấu hình AI_COACH_MODE=auto hoặc internal. "
        "Ở chế độ auto, hệ thống thử cloud AI trước; nếu lỗi quota/network sẽ fallback nội bộ. "
        "Ở chế độ internal, toàn bộ phản hồi AI được sinh từ luật nghiệp vụ và dữ liệu hành vi.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    h(doc, "3.1 Phân tích actor và use case", 2)
    tbl(
        doc,
        ["Actor", "Mục tiêu", "Use case chính"],
        [
            ["Khách", "Khám phá nền tảng", "Xem trang chủ, tìm kiếm khóa học"],
            ["Học viên", "Học tập và nhận hỗ trợ", "Đăng ký khóa, học bài, nhận AI coach"],
            ["Giảng viên", "Xây dựng nội dung", "Tạo/sửa/xóa khóa học và bài học"],
            ["Admin", "Quản trị vận hành", "Duyệt giao dịch, quản lý dữ liệu, dashboard"],
        ],
    )
    h(doc, "3.2 Đặc tả yêu cầu chức năng", 2)
    for x in [
        "FR01: Đăng ký, đăng nhập, đăng xuất.",
        "FR02: Tìm kiếm, lọc và xem danh sách khóa học.",
        "FR03: Đăng ký khóa học và thanh toán bằng ví nội bộ.",
        "FR04: Học bài, đánh dấu hoàn thành, cập nhật tiến độ.",
        "FR05: API AI cá nhân hóa trả profile + coach.",
        "FR06: Nạp/rút tiền theo cơ chế yêu cầu và duyệt.",
        "FR07: Admin từ chối rút tiền và hoàn tiền tự động.",
    ]:
        b(doc, x)
    h(doc, "3.3 Đặc tả yêu cầu phi chức năng", 2)
    for x in [
        "NFR01: Giao diện dễ dùng, trực quan, responsive.",
        "NFR02: Bảo mật route theo session và quyền truy cập.",
        "NFR03: Độ sẵn sàng cao nhờ cơ chế fallback AI.",
        "NFR04: Dễ bảo trì nhờ tách service layer.",
        "NFR05: Hiệu năng đáp ứng yêu cầu demo thực tế.",
    ]:
        b(doc, x)
    h(doc, "3.4 Thiết kế cơ sở dữ liệu", 2)
    tbl(
        doc,
        ["Bảng", "Mục đích"],
        [
            ["users", "Thông tin tài khoản, quyền, ví, ngân hàng"],
            ["courses", "Thông tin khóa học, mức độ, học phí"],
            ["lessons", "Nội dung bài học theo khóa"],
            ["enrollments", "Liên kết user-khóa và tiến độ tổng"],
            ["lesson_progress", "Theo dõi hoàn thành theo từng bài"],
            ["wallet_transactions", "Lịch sử nạp/rút/mua/hoàn tiền"],
            ["deposit_requests", "Yêu cầu nạp tiền chờ admin duyệt"],
            ["delete_requests", "Yêu cầu xóa tài khoản"],
        ],
    )
    h(doc, "3.5 Thiết kế luồng AI cá nhân hóa", 2)
    for x in [
        "B1: Thu thập dữ liệu hành vi (enrollments, lesson_progress).",
        "B2: Tính các chỉ số tiến độ và mức độ hoạt động.",
        "B3: Chấm risk_score và phân loại risk_level.",
        "B4: Sinh next_lessons và recommendations.",
        "B5: Tạo coach_message theo mode AI hiện tại.",
        "B6: Trả dữ liệu qua /api/ai/personalization cho frontend.",
    ]:
        n(doc, x)
    p(doc, "Hình 3.1. Sơ đồ use case tổng quan hệ thống.")
    p(doc, "Hình 3.2. Sơ đồ kiến trúc triển khai hệ thống.")
    p(doc, "Hình 3.3. Sơ đồ ER rút gọn cơ sở dữ liệu.")
    p(doc, "Hình 3.4. Sơ đồ luồng AI cá nhân hóa.")

    page_break(doc)
    h(doc, "CHƯƠNG 4. TRIỂN KHAI VÀ XÂY DỰNG CHỨC NĂNG", 1)
    h(doc, "4.1 Môi trường phát triển", 2)
    for x in [
        "Hệ điều hành: Windows 10.",
        "Python 3.10+.",
        "Flask + sqlite3 + python-docx.",
        "Trình duyệt kiểm thử: Chrome/Cốc Cốc.",
    ]:
        b(doc, x)
    h(doc, "4.2 Cấu trúc mã nguồn", 2)
    for x in [
        "app.py: route layer và integration.",
        "services/ai_personalization.py: toàn bộ logic AI.",
        "templates/: giao diện server-side render.",
        "static/js, static/css: xử lý client và trình bày.",
        "database.sql: schema và dữ liệu mẫu.",
    ]:
        b(doc, x)
    h(doc, "4.3 API tiêu biểu", 2)
    p(doc, "Bảng 4.1. Danh sách API tiêu biểu.")
    tbl(
        doc,
        ["API", "Method", "Input", "Output"],
        [
            ["/api/ai/personalization", "GET", "session user_id", "profile + coach"],
            ["/enroll/<id>", "POST", "course_id", "success + first_lesson"],
            ["/mark-lesson-complete", "POST", "lesson_id", "progress_pct"],
            ["/wallet/withdraw-request", "POST", "amount", "pending withdrawal"],
            ["/admin/withdrawals/reject/<id>", "POST", "note", "refund transaction"],
        ],
    )
    h(doc, "4.4 Cơ chế hoàn tiền khi từ chối rút", 2)
    p(
        doc,
        "Khi admin từ chối yêu cầu rút tiền ở trạng thái pending, hệ thống sẽ tự động cộng lại số dư ví cho người dùng, "
        "đánh dấu giao dịch rút là failed và ghi thêm một giao dịch refund để đảm bảo tính minh bạch.",
    )
    h(doc, "4.5 Triển khai AI mode trên UI", 2)
    p(
        doc,
        "Tab AI hiển thị rõ mode đang hoạt động: AI Coach (Gemini Mode) hoặc AI Coach (Internal/Free Mode). "
        "Nút Làm mới phản hồi hỗ trợ cập nhật realtime và auto-refresh khi mở tab AI.",
    )

    page_break(doc)
    h(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    h(doc, "5.1 Kịch bản kiểm thử chính", 2)
    p(doc, "Bảng 5.1. Danh sách test case chính.")
    tbl(
        doc,
        ["ID", "Mô tả kiểm thử", "Kết quả mong đợi"],
        [
            ["TC01", "Đăng ký tài khoản mới", "Tạo user thành công"],
            ["TC02", "Đăng ký khóa học", "Enrollment hợp lệ, cập nhật dữ liệu"],
            ["TC03", "Đánh dấu hoàn thành bài", "Tăng tiến độ khóa học"],
            ["TC04", "AI internal mode", "coach_source = rule_engine"],
            ["TC05", "AI auto mode khi cloud lỗi", "Fallback nội bộ hoạt động"],
            ["TC06", "Admin từ chối rút tiền", "Hoàn tiền + tạo refund transaction"],
        ],
    )
    h(doc, "5.2 Đánh giá kết quả đạt được", 2)
    b(doc, "Hệ thống vận hành ổn định ở các luồng chính.")
    b(doc, "AI cá nhân hóa hoạt động tốt trong cả cloud mode và internal mode.")
    b(doc, "Luồng ví và giao dịch minh bạch, có cơ chế hoàn tiền rõ ràng.")
    b(doc, "Code được tách lớp service, thuận lợi cho bảo trì và mở rộng.")
    h(doc, "5.3 Hạn chế hiện tại", 2)
    b(doc, "Chưa có bộ kiểm thử tự động toàn diện.")
    b(doc, "Chưa tối ưu cho tải lớn và triển khai phân tán.")
    b(doc, "Chưa tích hợp mobile native.")

    page_break(doc)
    h(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    h(doc, "6.1 Kết luận", 2)
    p(
        doc,
        f"Đề tài \"{TITLE}\" đã đạt được mục tiêu xây dựng một nền tảng LMS thông minh có khả năng cá nhân hóa học tập và theo dõi tiến độ. "
        "Giải pháp AI Hybrid là điểm mạnh, giúp hệ thống cân bằng giữa chất lượng và chi phí vận hành.",
    )
    h(doc, "6.2 Hướng phát triển", 2)
    for x in [
        "Bổ sung bộ kiểm thử tự động CI/CD.",
        "Nâng cấp adaptive learning dựa trên kết quả bài tập.",
        "Phát triển dashboard phân tích nâng cao cho giảng viên và admin.",
        "Tích hợp thông báo email/app theo sự kiện học tập.",
        "Mở rộng kiến trúc để phục vụ tải lớn hơn.",
    ]:
        n(doc, x)

    page_break(doc)
    h(doc, "PHỤ LỤC SƠ ĐỒ VÀ TÀI LIỆU KỸ THUẬT", 1)
    diagram_files = [
        ("Phụ lục A. Sơ đồ Use Case", DOCS_DIR / "so-do-usecase.mmd"),
        ("Phụ lục B. Sơ đồ kiến trúc hệ thống", DOCS_DIR / "so-do-kien-truc-he-thong.mmd"),
        ("Phụ lục C. Sơ đồ ER rút gọn", DOCS_DIR / "so-do-er-rut-gon.mmd"),
        ("Phụ lục D. Sơ đồ luồng AI cá nhân hóa", DOCS_DIR / "so-do-luong-ai-ca-nhan-hoa.mmd"),
    ]
    for title, fp in diagram_files:
        h(doc, title, 2)
        p(doc, f"Nguồn sơ đồ: {fp.name}")
        if fp.exists():
            add_code_block(doc, fp.read_text(encoding="utf-8"))
        else:
            p(doc, "Không tìm thấy file sơ đồ.")


def main():
    doc = Document()
    cover(doc)
    review_pages(doc)
    content(doc)
    apply_global_format(doc)
    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
