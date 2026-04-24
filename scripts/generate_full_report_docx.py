from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "TEMPT_EduConnect_Full.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str):
    return doc.add_paragraph(text)


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str):
    doc.add_paragraph(text, style="List Number")


def add_table_simple(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def add_page_break(doc: Document):
    doc.add_page_break()


def build_cover(doc: Document):
    p = add_para(doc, "TRUONG DAI HOC ........................................")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_para(doc, "KHOA ......................................................")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_para(doc, "")
    p = add_para(doc, "BAO CAO DO AN TOT NGHIEP")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(20)

    p = add_para(doc, "DE TAI")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    p = add_para(doc, "HE THONG QUAN LY HOC TAP THONG MINH EDUCONNECT")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    add_para(doc, "")
    add_para(doc, "Sinh vien thuc hien: ............................................................")
    add_para(doc, "MSSV: ...........................................................................")
    add_para(doc, "Lop: .............................................................................")
    add_para(doc, "Giang vien huong dan: ............................................................")
    add_para(doc, "Nganh: Cong nghe thong tin")
    add_para(doc, "")
    p = add_para(doc, "Binh Duong, nam 2026")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def build_sections(doc: Document):
    add_page_break(doc)
    add_heading(doc, "LOI CAM ON", 1)
    add_para(
        doc,
        "Em xin chan thanh cam on quy Thay/Cô khoa Cong nghe thong tin da trang bi kien thuc nen tang "
        "va huong dan em trong suot qua trinh hoc tap. Em xin gui loi cam on den Giang vien huong dan "
        "da tan tinh gop y de em hoan thien de tai nay.",
    )

    add_heading(doc, "NHAN XET CUA GIANG VIEN HUONG DAN", 1)
    add_para(doc, "........................................................................................................")
    add_para(doc, "........................................................................................................")

    add_heading(doc, "TOM TAT DE TAI", 1)
    add_para(
        doc,
        "De tai trinh bay viec xay dung he thong quan ly hoc tap truc tuyen co tich hop AI ca nhan hoa. "
        "He thong ho tro nguoi hoc theo doi tien do, du doan rui ro bo hoc, goi y bai hoc tiep theo va "
        "goi y khoa hoc phu hop. Nen tang duoc phat trien bang Flask + SQLite, frontend HTML/CSS/JS.",
    )

    add_page_break(doc)
    add_heading(doc, "MUC LUC (THU CONG)", 1)
    toc_items = [
        "Chuong 1. Tong quan de tai",
        "Chuong 2. Co so ly thuyet va cong nghe",
        "Chuong 3. Phan tich va thiet ke he thong",
        "Chuong 4. Trien khai he thong",
        "Chuong 5. Kiem thu va danh gia",
        "Chuong 6. Ket luan va huong phat trien",
        "Phu luc",
    ]
    for item in toc_items:
        add_number(doc, item)

    add_page_break(doc)
    add_heading(doc, "CHUONG 1. TONG QUAN DE TAI", 1)
    add_heading(doc, "1.1 Ly do chon de tai", 2)
    add_para(
        doc,
        "Trong boi canh chuyen doi so giao duc, nhu cau hoc truc tuyen ngay cang cao. "
        "Tuy nhien, nhieu nen tang hien tai chua toi uu kha nang ca nhan hoa theo nang luc va hanh vi hoc. "
        "Do do, de tai huong den viec xay dung he thong hoc tap thong minh, co kha nang phan tich va dua ra goi y phu hop.",
    )
    add_heading(doc, "1.2 Muc tieu de tai", 2)
    goals = [
        "Xay dung nen tang LMS co day du vai tro hoc vien, giang vien, quan tri vien.",
        "Quan ly khoa hoc, bai hoc, tai lieu va bai tap.",
        "Theo doi tien do hoc tap theo tung bai hoc.",
        "Ung dung AI de goi y lo trinh hoc ca nhan hoa.",
        "Xay dung co che fallback AI noi bo de tiet kiem chi phi.",
    ]
    for g in goals:
        add_bullet(doc, g)
    add_heading(doc, "1.3 Pham vi de tai", 2)
    add_para(
        doc,
        "De tai tap trung vao he thong web don vi monolithic de phuc vu demo do an. "
        "Cac tinh nang mobile native va tich hop thanh toan thuc te khong nam trong pham vi ban dau.",
    )

    add_page_break(doc)
    add_heading(doc, "CHUONG 2. CO SO LY THUYET VA CONG NGHE", 1)
    add_heading(doc, "2.1 Kien thuc nen tang", 2)
    add_bullet(doc, "Kien truc web client-server.")
    add_bullet(doc, "Mo hinh MVC tinh gon trong Flask.")
    add_bullet(doc, "Quan tri du lieu giao dich voi SQLite.")
    add_bullet(doc, "Nguyen ly recommendation va risk scoring.")
    add_heading(doc, "2.2 Cong nghe su dung", 2)
    add_table_simple(
        doc,
        ["Thanh phan", "Cong nghe", "Muc dich"],
        [
            ["Backend", "Python Flask", "Xu ly route, business logic, API"],
            ["Database", "SQLite", "Luu tru users, courses, progress, wallet"],
            ["Frontend", "HTML/CSS/JS", "Giao dien va tuong tac nguoi dung"],
            ["AI Cloud", "Gemini API", "Sinh coach message nang cao"],
            ["AI Noi bo", "Rule-based Engine", "Fallback mien phi, on dinh"],
        ],
    )
    add_heading(doc, "2.3 Uu diem va han che", 2)
    add_para(
        doc,
        "Giai phap co uu diem de trien khai nhanh, chi phi thap, de demo va de bao tri. "
        "Han che chinh la chua toi uu cho tai luong lon va chua trien khai vi mo microservices.",
    )

    add_page_break(doc)
    add_heading(doc, "CHUONG 3. PHAN TICH VA THIET KE HE THONG", 1)
    add_heading(doc, "3.1 Phan tich yeu cau chuc nang", 2)
    functional = [
        "Dang ky, dang nhap, quan ly tai khoan.",
        "Tim kiem, loc, dang ky khoa hoc.",
        "Hoc bai hoc, danh dau hoan thanh.",
        "Quan ly khoa hoc cho giang vien.",
        "Quan tri nguoi dung va giao dich cho admin.",
        "AI ca nhan hoa va goi y lo trinh.",
    ]
    for f in functional:
        add_bullet(doc, f)
    add_heading(doc, "3.2 Yeu cau phi chuc nang", 2)
    non_func = [
        "He thong de su dung, giao dien ro rang, responsive.",
        "Bao mat session va phan quyen theo vai tro.",
        "Do tin cay cao nhat la fallback khi cloud AI loi.",
        "Thoi gian phan hoi API dat muc chap nhan duoc.",
    ]
    for n in non_func:
        add_bullet(doc, n)

    add_heading(doc, "3.3 Thiet ke co so du lieu", 2)
    add_para(doc, "Cac bang du lieu chinh trong he thong:")
    add_table_simple(
        doc,
        ["Bang", "Muc dich chinh"],
        [
            ["users", "Thong tin nguoi dung, quyen, vi, thong tin ngan hang"],
            ["courses", "Thong tin khoa hoc, gia, level, danh muc"],
            ["lessons", "Danh sach bai hoc cua tung khoa"],
            ["enrollments", "Quan he user-khoa hoc va tien do tong"],
            ["lesson_progress", "Tien do theo bai hoc"],
            ["wallet_transactions", "Nhat ky nap/rut/mua/hoan tien"],
            ["deposit_requests", "Yeu cau nap tien cho admin duyet"],
            ["delete_requests", "Yeu cau xoa tai khoan"],
        ],
    )

    add_heading(doc, "3.4 Thiet ke luong AI ca nhan hoa", 2)
    add_number(doc, "Thu thap du lieu hanh vi hoc tap cua nguoi dung.")
    add_number(doc, "Tinh toan risk_score, risk_level, inactivity_days.")
    add_number(doc, "Goi y bai hoc tiep theo va khoa hoc de xuat.")
    add_number(doc, "Sinh coach message theo AI mode (auto/internal).")
    add_number(doc, "Tra ve ket qua qua API /api/ai/personalization.")

    add_page_break(doc)
    add_heading(doc, "CHUONG 4. TRIEN KHAI HE THONG", 1)
    add_heading(doc, "4.1 Cau truc thu muc", 2)
    add_para(
        doc,
        "Du an duoc to chuc theo kieu monolithic: app.py, templates/, static/, database.sql, services/. "
        "Trong do services/ chua cac module logic de giup code de doc va de bao tri.",
    )
    add_heading(doc, "4.2 Trien khai module AI", 2)
    add_bullet(doc, "services/ai_personalization.py: tinh toan profile va goi y.")
    add_bullet(doc, "AI_COACH_MODE=auto: uu tien cloud AI, loi thi fallback.")
    add_bullet(doc, "AI_COACH_MODE=internal: luon dung AI noi bo mien phi.")
    add_bullet(doc, "UI cap nhat nhan mode: Gemini Mode / Internal Free Mode.")
    add_heading(doc, "4.3 Trien khai quy trinh vi va giao dich", 2)
    add_para(
        doc,
        "He thong ho tro yeu cau rut tien cho admin duyet. Trong truong hop admin tu choi rut, "
        "so tien duoc hoan lai vao vi va ghi nhat ky giao dich loai refund de minh bach.",
    )

    add_page_break(doc)
    add_heading(doc, "CHUONG 5. KIEM THU VA DANH GIA", 1)
    add_heading(doc, "5.1 Kich ban kiem thu chinh", 2)
    add_table_simple(
        doc,
        ["Ma test", "Mo ta", "Ket qua mong doi"],
        [
            ["TC01", "Dang ky tai khoan moi", "Tao user thanh cong"],
            ["TC02", "Dang ky khoa hoc", "Them enrollment, tru vi neu co phi"],
            ["TC03", "Danh dau hoan thanh bai", "Cap nhat lesson_progress"],
            ["TC04", "AI mode internal", "coach_source = rule_engine"],
            ["TC05", "AI mode auto + cloud loi", "Fallback noi bo, he thong van chay"],
            ["TC06", "Admin tu choi rut tien", "Hoan tien + tao refund transaction"],
        ],
    )
    add_heading(doc, "5.2 Danh gia ket qua", 2)
    add_para(
        doc,
        "Ket qua thuc nghiem cho thay he thong dap ung duoc cac muc tieu chinh dat ra. "
        "AI noi bo van hoat dong tot khi khong co cloud quota, giup de tai co tinh ung dung thuc te cao.",
    )
    add_heading(doc, "5.3 Danh gia hieu nang va on dinh", 2)
    add_bullet(doc, "Thoi gian phan hoi trang trong muc chap nhan cho demo do an.")
    add_bullet(doc, "Fallback giup giam nguy co gian doan dich vu.")
    add_bullet(doc, "Code da duoc tach module de de bao tri va mo rong.")

    add_page_break(doc)
    add_heading(doc, "CHUONG 6. KET LUAN VA HUONG PHAT TRIEN", 1)
    add_heading(doc, "6.1 Ket luan", 2)
    add_para(
        doc,
        "De tai da xay dung thanh cong he thong quan ly hoc tap thong minh EduConnect "
        "voi day du chuc nang cot loi va bo AI ca nhan hoa dua tren du lieu hanh vi nguoi hoc.",
    )
    add_heading(doc, "6.2 Huong phat trien", 2)
    future = [
        "Tich hop cache va queue de toi uu khi tai luong lon.",
        "Bo sung bo test tu dong (unit/integration/e2e).",
        "Mo rong AI theo adaptive sequencing nang cao.",
        "Tich hop thong bao email/app theo su kien hoc tap.",
        "Bo sung dashboard phan tich hoc tap nang cao cho admin.",
    ]
    for f in future:
        add_bullet(doc, f)

    add_page_break(doc)
    add_heading(doc, "PHU LUC", 1)
    add_heading(doc, "Phu luc A - Danh sach API tieu bieu", 2)
    add_table_simple(
        doc,
        ["Method", "Endpoint", "Mo ta"],
        [
            ["GET", "/khoa-hoc", "Danh sach khoa hoc va loc"],
            ["POST", "/enroll/<id>", "Dang ky khoa hoc"],
            ["POST", "/mark-lesson-complete", "Danh dau hoan thanh bai hoc"],
            ["GET", "/api/ai/personalization", "Lay goi y AI ca nhan hoa"],
            ["POST", "/wallet/withdraw-request", "Gui yeu cau rut tien"],
            ["POST", "/admin/withdrawals/reject/<id>", "Tu choi rut va hoan tien"],
        ],
    )
    add_heading(doc, "Phu luc B - So do", 2)
    add_bullet(doc, "so-do-usecase.mmd")
    add_bullet(doc, "so-do-kien-truc-he-thong.mmd")
    add_bullet(doc, "so-do-er-rut-gon.mmd")
    add_bullet(doc, "so-do-luong-ai-ca-nhan-hoa.mmd")


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    build_cover(doc)
    build_sections(doc)
    doc.save(OUT_FILE)
    print(f"Generated full report: {OUT_FILE}")


if __name__ == "__main__":
    main()
