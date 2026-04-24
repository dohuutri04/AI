from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUT_FILE = DOCS_DIR / "BaoCao_DoAn_EduConnect.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_diagram_appendix(doc: Document, title: str, file_path: Path):
    add_heading(doc, title, level=2)
    doc.add_paragraph(f"Nguon: {file_path.name}")
    if not file_path.exists():
        doc.add_paragraph("Khong tim thay file so do.")
        return
    content = file_path.read_text(encoding="utf-8")
    p = doc.add_paragraph(content)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def main():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    add_heading(doc, "BAO CAO DO AN NGANH", level=0)
    doc.add_paragraph("De tai: He thong quan ly hoc tap thong minh EduConnect")
    doc.add_paragraph("Sinh vien: ....................................................")
    doc.add_paragraph("MSSV: ........................................................")
    doc.add_paragraph("Giang vien huong dan: ........................................")
    doc.add_paragraph("Nam hoc: 2025 - 2026")

    add_heading(doc, "1. Gioi thieu de tai", level=1)
    doc.add_paragraph(
        "EduConnect la nen tang hoc truc tuyen su dung Flask + SQLite, "
        "tap trung vao quan ly khoa hoc, theo doi tien do hoc tap va ca nhan hoa bang AI."
    )

    add_heading(doc, "2. Muc tieu de tai", level=1)
    add_bullet(doc, "Xay dung he thong hoc truc tuyen co day du vai tro hoc vien, giang vien va admin.")
    add_bullet(doc, "Theo doi tien do hoc tap va giam sat hanh vi hoc theo thoi gian thuc.")
    add_bullet(doc, "Ung dung AI de ca nhan hoa lo trinh hoc va canh bao nguy co bo hoc.")
    add_bullet(doc, "Tich hop vi noi bo va quy trinh nap/rut co duyet.")

    add_heading(doc, "3. Chuc nang trong tam", level=1)
    add_heading(doc, "3.1 Chuc nang he thong", level=2)
    add_bullet(doc, "Dang ky, dang nhap, quan ly tai khoan nguoi dung.")
    add_bullet(doc, "Tim kiem, loc, dang ky khoa hoc.")
    add_bullet(doc, "Hoc bai hoc, danh dau hoan thanh, cap nhat tien do.")
    add_bullet(doc, "Quan ly khoa hoc cho giang vien.")
    add_bullet(doc, "Dashboard admin: quan ly nguoi dung, danh muc, lien he.")
    add_bullet(doc, "Nap/rut vi co duyet, tu choi rut co hoan tien tu dong.")

    add_heading(doc, "3.2 Chuc nang AI", level=2)
    add_bullet(doc, "Tinh toan hoc vien profile: avg progress, completed, stalled, inactivity.")
    add_bullet(doc, "Dropout risk scoring: risk_score (0-100) va risk_level.")
    add_bullet(doc, "Goi y bai hoc tiep theo theo tung khoa dang hoc.")
    add_bullet(doc, "Goi y khoa hoc ca nhan hoa theo muc do va chu de.")
    add_bullet(doc, "AI Coach sinh phan hoi bang Gemini, co fallback rule-based.")
    add_bullet(doc, "Lam moi phan hoi AI realtime thu cong va tu dong.")

    add_heading(doc, "4. Cong nghe su dung", level=1)
    add_bullet(doc, "Backend: Python Flask.")
    add_bullet(doc, "Database: SQLite.")
    add_bullet(doc, "Frontend: HTML, CSS, JavaScript.")
    add_bullet(doc, "AI: Rule-based engine + Gemini API.")

    add_heading(doc, "5. Thiet ke co so du lieu (tom tat)", level=1)
    add_bullet(doc, "users, courses, lessons, enrollments, lesson_progress.")
    add_bullet(doc, "wallet_transactions, deposit_requests, delete_requests.")
    add_bullet(doc, "categories, contacts, reviews.")

    add_heading(doc, "6. Quy trinh xu ly AI ca nhan hoa", level=1)
    doc.add_paragraph(
        "Khi nguoi dung mo tab AI, frontend goi API /api/ai/personalization. "
        "Backend phan tich du lieu hoc tap, tinh risk score, sinh goi y bai hoc va "
        "goi y khoa hoc. Neu co GEMINI_API_KEY thi goi Gemini de tao coach message; "
        "neu khong he thong fallback ve noi dung rule-based."
    )

    add_heading(doc, "7. Danh sach so do kem theo", level=1)
    add_bullet(doc, "so-do-usecase.mmd")
    add_bullet(doc, "so-do-kien-truc-he-thong.mmd")
    add_bullet(doc, "so-do-er-rut-gon.mmd")
    add_bullet(doc, "so-do-luong-ai-ca-nhan-hoa.mmd")
    doc.add_paragraph(
        "Ban co the mo cac file .mmd bang Mermaid Editor de xuat PNG/SVG dua vao slide/bao cao."
    )

    add_heading(doc, "8. Ket luan va huong phat trien", level=1)
    add_bullet(doc, "Hoan thien CSRF, OTP reset password va test tu dong.")
    add_bullet(doc, "Mo rong AI voi learning outcome prediction va adaptive sequencing.")
    add_bullet(doc, "Bo sung thong bao email/app notification khi duyet giao dich.")

    add_heading(doc, "Phu luc so do (gop chung trong file DOCX)", level=1)
    add_diagram_appendix(doc, "A. So do Use Case", DOCS_DIR / "so-do-usecase.mmd")
    add_diagram_appendix(doc, "B. So do kien truc he thong", DOCS_DIR / "so-do-kien-truc-he-thong.mmd")
    add_diagram_appendix(doc, "C. So do ER rut gon", DOCS_DIR / "so-do-er-rut-gon.mmd")
    add_diagram_appendix(doc, "D. So do luong AI ca nhan hoa", DOCS_DIR / "so-do-luong-ai-ca-nhan-hoa.mmd")

    doc.save(OUT_FILE)
    print(f"Generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
